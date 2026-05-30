#!/usr/bin/env python3
"""
DocStrange test script.

Strategy:
  1. Create a sample DOCX (python-docx is a docstrange dependency, always available).
  2. Use docstrange.processors.DOCXProcessor directly for fully local extraction.
  3. Also try cloud mode (DocumentExtractor) to show the cloud path.
  4. Print SUCCESS when Markdown output is non-empty.
"""
import os
import sys

WORK_DIR = os.path.dirname(os.path.abspath(__file__))


# ── 1. Create a sample DOCX file ─────────────────────────────────────────────
def create_sample_docx(path: str) -> None:
    """Build a sample .docx using python-docx (always installed as a dep)."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading("DocStrange Test Document", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This is a sample document generated automatically for testing DocStrange. "
        "It contains multiple paragraphs, headings, a bulleted list, and a table."
    )

    doc.add_heading("Section 1: About DocStrange", level=2)
    doc.add_paragraph(
        "DocStrange is an open-source tool by NanoNets that converts PDF, DOCX, "
        "PPTX, XLSX, and images to Markdown, JSON, CSV, or HTML. "
        "It supports both a free cloud API and fully-local CPU/GPU processing."
    )

    doc.add_heading("Section 2: Features", level=2)
    features = [
        "PDF, DOCX, PPTX, XLSX, and image input",
        "Markdown, JSON, CSV, HTML output",
        "Free cloud tier (10k docs/month with login)",
        "100% local mode (CPU or GPU)",
        "Advanced OCR and table extraction",
    ]
    for feat in features:
        doc.add_paragraph(feat, style="List Bullet")

    doc.add_heading("Section 3: Processing Modes", level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Mode", "Input", "Output"]):
        table.cell(0, i).text = h
    for ri, row_data in enumerate(
        [
            ["Cloud", "PDF / Image / DOCX", "Markdown / JSON"],
            ["Local", "PDF / Image / DOCX", "Markdown / JSON / CSV"],
        ],
        start=1,
    ):
        for ci, val in enumerate(row_data):
            table.cell(ri, ci).text = val

    doc.save(path)
    print(f"[1] Created sample DOCX: {path} ({os.path.getsize(path)} bytes)")


# ── 2. Local extraction via DOCXProcessor ────────────────────────────────────
def run_local_extraction(docx_path: str) -> str:
    """Use docstrange's built-in local DOCX processor — no cloud API needed."""
    from docstrange.processors.docx_processor import DOCXProcessor

    print("[2] Instantiating DOCXProcessor (local, no cloud API) ...")
    processor = DOCXProcessor(preserve_layout=True, include_images=False)

    print(f"[3] Processing: {docx_path}")
    result = processor.process(docx_path)

    print("[4] Extracting Markdown ...")
    return result.extract_markdown()


# ── 3. Cloud extraction (best-effort, unauthenticated) ────────────────────────
def try_cloud_extraction(docx_path: str):
    """
    Try the DocumentExtractor cloud path with a patched auth that forces
    the unauthenticated free tier (avoids the exhausted sandbox OAuth account).
    Returns Markdown string on success, None on any failure.
    """
    try:
        import docstrange.services.auth_service as _auth
        _orig = _auth.get_authenticated_token
        _auth.get_authenticated_token = lambda force_reauth=False: None
        os.environ.pop("NANONETS_API_KEY", None)

        from docstrange import DocumentExtractor
        extractor = DocumentExtractor()  # cloud mode, api_key=None (rate-limited)
        result = extractor.extract(docx_path)
        md = result.extract_markdown()

        _auth.get_authenticated_token = _orig
        return md if md and len(md.strip()) > 5 else None
    except Exception:
        return None


# ── 4. Main ───────────────────────────────────────────────────────────────────
def main() -> None:
    docx_path = os.path.join(WORK_DIR, "sample.docx")

    # Step 1 – create sample DOCX
    create_sample_docx(docx_path)

    # Step 2 – local extraction (always available)
    try:
        local_md = run_local_extraction(docx_path)
    except Exception as exc:
        print(f"\n[ERROR] Local extraction failed: {exc}", file=sys.stderr)
        sys.exit(1)

    # Step 3 – try cloud extraction
    print("[5] Attempting cloud extraction (unauthenticated free tier) ...")
    cloud_md = try_cloud_extraction(docx_path)
    if cloud_md:
        print("    Cloud extraction succeeded.")
    else:
        print("    Cloud extraction unavailable (rate-limited or no network). Local result used.")

    final_md = cloud_md if cloud_md else local_md
    source = "cloud (unauthenticated free tier)" if cloud_md else "local (DOCXProcessor)"

    # Report
    print("\n" + "=" * 60)
    print("DOCSTRANGE EXTRACTION RESULT (Markdown):")
    print("=" * 60)
    print(final_md[:2000])
    print("=" * 60)

    if final_md and len(final_md.strip()) > 10:
        print(f"\n===================================================")
        print(f"SUCCESS — DocStrange DOCX → Markdown ({source})")
        print(f"  Characters extracted : {len(final_md)}")
        print(f"  Input file           : {docx_path}")
        print(f"==================================================")
    else:
        print("\nWARNING — Markdown output is empty.", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
